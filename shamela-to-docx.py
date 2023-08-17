from urllib.parse import urlparse, urlunparse, ParseResult
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_last_page_url(url):
    try:
        # Send request and get response
        response = requests.get(url)
        response.raise_for_status()  # Check for request errors

        html_content = response.content

        # Parse HTML content using Beautiful Soup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Find the main <div> with class 'col-md-8'
        main_div = soup.find('div', class_='col-md-8')

        if main_div:
            # Find the second nested <div> inside the main <div>
            nested_div = main_div.find_all('div')[1]

            # Find the last <a> element inside the nested <div>
            last_a_inside_nested_div = nested_div.find_all('a')[-1]

            if last_a_inside_nested_div:
                last_page_url = last_a_inside_nested_div['href']
                return last_page_url
            else:
                print("No <a> element found inside the nested <div>.")
        else:
            print("Main <div> element not found.")
    except requests.exceptions.RequestException as e:
        print("Error:", e)
    
    return None

def extract_chapter_number_from_url(url):
    parsed_url = urlparse(url)
    path_segments = parsed_url.path.split('/')
    if len(path_segments) >= 4:
        return path_segments[3]
    else:
        return None


def extract_text_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        html_content = response.content
        soup = BeautifulSoup(html_content, 'html.parser')

        content_div = soup.find('div', class_='nass')
        
        if content_div:
            paragraphs = content_div.find_all('p')
            text = '\n'.join(paragraph.get_text() for paragraph in paragraphs)
            return text
        else:
            return None
    except requests.exceptions.RequestException as e:
        print("Error:", e)
        return None

def create_combined_docx(urls, output_filename='combined-output.docx'):
    doc = Document()

    for url in urls:
        text = extract_text_from_url(url)
        if text:
            paragraph = doc.add_paragraph(text)
            run = paragraph.runs[-1]  # Get the run of the last added paragraph
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save(output_filename)
    #print(f"Combined text from {len(urls)} URLs saved to '{output_filename}'")


def normalize_url(url):
    parsed_url = urlparse(url)
    if parsed_url.netloc != 'shamela.ws' or not parsed_url.path.startswith('/book/'):
        print("Error: URL should begin with 'https://shamela.ws/book/'")
        return None
    
    # Ensure the path ends with a '/'
    if not parsed_url.path.endswith('/'):
        parsed_url = parsed_url._replace(path=parsed_url.path + '/')

    path_segments = parsed_url.path.split('/')
    if len(path_segments) >= 4:
        identifier = path_segments[2]
        chapter_number = path_segments[3] if len(path_segments) >= 5 else '1'
        return urlunparse(ParseResult(scheme=parsed_url.scheme, netloc=parsed_url.netloc,
                                      path=f"/book/{identifier}/{chapter_number}", params='', query='', fragment=''))
    else:
        print("Error: Incorrect URL format.")
        return None



def process_multiple_pages(start_url, output_filename='combined-output.docx'):
    normalized_url = normalize_url(start_url)
    if normalized_url is None:
        return
    
    last_page_url = get_last_page_url(normalized_url)
    if last_page_url is None:
        print("Unable to determine the last page URL.")
        return
    
    book_identifier = urlparse(normalized_url).path.split('/')[2]

    chapter_start = int(extract_chapter_number_from_url(normalized_url))
    chapter_end = int(extract_chapter_number_from_url(last_page_url))

    urls = []
    for chapter_number in range(chapter_start, chapter_end + 1):
        chapter_url = f"https://shamela.ws/book/{book_identifier}/{chapter_number}"
        urls.append(chapter_url)

    stripped_url = f"https://shamela.ws/book/{book_identifier}/"
    doc = Document()  # Create a new DOCX document

    # Add the first page content using the provided code
    response = requests.get(stripped_url)
    soup = BeautifulSoup(response.content, "html.parser")
    betaka_div = soup.find("div", class_="betaka-index")
    nass_div = soup.find("div", class_="nass")
    h1_element = soup.find("h1")
    font_name = "Calibri"

    def add_separator():
        doc.add_paragraph("-------------").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    if h1_element:
        h1_text = h1_element.get_text()
        h1_paragraph = doc.add_paragraph(h1_text)
        h1_run = h1_paragraph.runs[0]
        h1_run.bold = True
        h1_run.font.size = Pt(21)
        h1_run.font.name = font_name
        add_separator()

    for element in nass_div.contents:
        if element.name == "h3":
            h3_text = element.get_text()
            h3_paragraph = doc.add_paragraph(h3_text)
            h3_run = h3_paragraph.runs[0]
            h3_run.bold = True
            h3_run.font.size = Pt(15)
            h3_run.font.name = font_name
        elif element.name is None and element.strip():
            formatted_line = element.replace('(', ']').replace(')', '[')
            formatted_line = formatted_line.replace('«', ']').replace('»', '[')
            nass_paragraph = doc.add_paragraph(formatted_line)
            nass_run = nass_paragraph.runs[0]
            nass_run.font.size = Pt(15)
            nass_run.font.name = font_name

    add_separator()

    for h4 in betaka_div.find_all("h4"):
        h4_text = h4.get_text()
        h4_paragraph = doc.add_paragraph(h4_text)
        h4_run = h4_paragraph.runs[0]
        h4_run.bold = True
        h4_run.font.size = Pt(15)
        h4_run.font.name = font_name

    for li in betaka_div.find_all("li"):
        li_text = li.get_text().strip('- ')
        li_paragraph = doc.add_paragraph(li_text)
        li_run = li_paragraph.runs[0]
        li_run.font.size = Pt(15)
        li_run.font.name = font_name

    doc.add_page_break()
    # Continue with the rest of the code
    for url in urls:
        text = extract_text_from_url(url)
        if text:
            paragraph = doc.add_paragraph(text)
            run = paragraph.runs[-1]  # Get the run of the last added paragraph
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    for paragraph in doc.paragraphs:
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save(output_filename)

    print(f"Combined text from Chapter {chapter_start} to {chapter_end} saved to '{output_filename}'")

# Input the starting URL
start_url = input("Enter the first URL: ")
#start_url = 'https://shamela.ws/book/6388'
# Call the function to process multiple pages and create the combined DOCX document
process_multiple_pages(start_url)
print(output_filename)
